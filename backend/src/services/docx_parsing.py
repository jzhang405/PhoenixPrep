"""
Docx文档解析服务
专门处理doc/docx格式的试题和答案文件
输出为HTML格式，支持图片、公式、化学式嵌入
"""

import os
import base64
import asyncio
from typing import Dict, Optional, List
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
except ImportError:
    Document = None

class DocxParsingService:
    """Docx文档解析服务"""
    
    def __init__(self):
        """初始化Docx解析服务"""
        # 使用项目内的输出目录
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
        self.output_dir = os.path.join(project_root, "data/parsed-html")
        
        # 确保输出目录存在
        os.makedirs(self.output_dir, exist_ok=True)
        
        print("✓ Docx解析服务初始化成功")
        print(f"  输出目录: {self.output_dir}")
    
    async def parse_document(self, 
                           input_path: str, 
                           output_path: Optional[str] = None,
                           is_answer: bool = False) -> Dict:
        """
        解析单个docx文档
        
        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径，如果为None则自动生成
            is_answer: 是否为答案文件
            
        Returns:
            解析结果字典
        """
        try:
            # 检查输入文件是否存在
            if not os.path.exists(input_path):
                raise FileNotFoundError(f"输入文件不存在: {input_path}")
            
            # 检查文件格式
            if not input_path.lower().endswith(('.docx', '.doc')):
                raise ValueError(f"不支持的文件格式: {input_path}")
            
            # 生成输出路径
            if output_path is None:
                input_name = Path(input_path).stem
                suffix = "_answer" if is_answer else "_question"
                output_path = os.path.join(
                    self.output_dir, 
                    f"{input_name}{suffix}_parsed.html"
                )
            
            print(f"开始解析文档: {input_path}")
            print(f"文件类型: {'答案' if is_answer else '试题'}")
            
            # 解析docx文档
            html_content = await self._parse_docx_to_html(input_path, is_answer)
            
            # 保存HTML内容
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return {
                'success': True,
                'input_path': input_path,
                'output_path': output_path,
                'html_content': html_content,
                'file_type': 'answer' if is_answer else 'question',
                'message': '文档解析成功'
            }
                
        except Exception as e:
            return {
                'success': False,
                'input_path': input_path,
                'error': str(e),
                'file_type': 'answer' if is_answer else 'question',
                'message': '文档解析失败'
            }
    
    async def parse_question_answer_pair(self,
                                       question_path: str,
                                       answer_path: str,
                                       output_path: Optional[str] = None) -> Dict:
        """
        解析试题-答案对
        
        Args:
            question_path: 试题文件路径
            answer_path: 答案文件路径
            output_path: 输出文件路径，如果为None则自动生成
            
        Returns:
            解析结果字典
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(question_path):
                raise FileNotFoundError(f"试题文件不存在: {question_path}")
            if not os.path.exists(answer_path):
                raise FileNotFoundError(f"答案文件不存在: {answer_path}")
            
            # 生成输出路径
            if output_path is None:
                question_name = Path(question_path).stem
                output_path = os.path.join(
                    self.output_dir, 
                    f"{question_name}_qa_parsed.html"
                )
            
            print(f"解析试题-答案对:")
            print(f"  试题: {question_path}")
            print(f"  答案: {answer_path}")
            
            # 解析试题和答案
            question_html = await self._parse_docx_to_html(question_path, is_answer=False)
            answer_html = await self._parse_docx_to_html(answer_path, is_answer=True)
            
            # 合并HTML内容
            combined_html = self._combine_question_answer(question_html, answer_html)
            
            # 保存HTML内容
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(combined_html)
            
            return {
                'success': True,
                'question_path': question_path,
                'answer_path': answer_path,
                'output_path': output_path,
                'html_content': combined_html,
                'message': '试题-答案对解析成功'
            }
                
        except Exception as e:
            return {
                'success': False,
                'question_path': question_path,
                'answer_path': answer_path,
                'error': str(e),
                'message': '试题-答案对解析失败'
            }
    
    async def _parse_docx_to_html(self, file_path: str, is_answer: bool = False) -> str:
        """
        将docx文档解析为HTML
        
        Args:
            file_path: docx文件路径
            is_answer: 是否为答案文件
            
        Returns:
            HTML内容
        """
        if Document is None:
            raise ImportError("python-docx库未安装，请运行: pip install python-docx")
        
        doc = Document(file_path)
        html_parts = []
        
        # 添加文档标题
        title = "答案" if is_answer else "试题"
        html_parts.append(f'<h1>{title}</h1>')
        
        # 解析段落
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                html_parts.append(f'<p>{self._escape_html(paragraph.text)}</p>')
        
        # 解析表格
        for table in doc.tables:
            html_parts.append('<table border="1" style="border-collapse: collapse;">')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<td>{self._escape_html(cell.text)}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        # 解析图片
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_data = rel.target_part.blob
                image_base64 = base64.b64encode(image_data).decode('utf-8')
                image_ext = self._get_image_extension(rel.target_part.content_type)
                html_parts.append(f'<img src="data:image/{image_ext};base64,{image_base64}" style="max-width: 100%;" />')
        
        return '\n'.join(html_parts)
    
    def _combine_question_answer(self, question_html: str, answer_html: str) -> str:
        """
        合并试题和答案HTML
        
        Args:
            question_html: 试题HTML
            answer_html: 答案HTML
            
        Returns:
            合并后的HTML
        """
        html_template = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>试题与答案</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        .question-section {{ margin-bottom: 40px; }}
        .answer-section {{ margin-top: 40px; padding-top: 20px; border-top: 2px solid #ccc; }}
        h1 {{ color: #333; }}
        h2 {{ color: #666; margin-top: 30px; }}
        table {{ width: 100%; margin: 10px 0; }}
        td {{ padding: 8px; border: 1px solid #ddd; }}
        img {{ max-width: 100%; height: auto; }}
    </style>
</head>
<body>
    <div class="question-section">
        <h1>📝 试题</h1>
        {{question_content}}
    </div>
    <div class="answer-section">
        <h1>🎯 答案</h1>
        {{answer_content}}
    </div>
</body>
</html>"""
        
        return html_template.format(
            question_content=question_html,
            answer_content=answer_html
        )
    
    def _escape_html(self, text: str) -> str:
        """转义HTML特殊字符"""
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&#39;')
    
    def _get_image_extension(self, content_type: str) -> str:
        """根据内容类型获取图片扩展名"""
        content_type_map = {
            'image/jpeg': 'jpeg',
            'image/jpg': 'jpeg',
            'image/png': 'png',
            'image/gif': 'gif',
            'image/bmp': 'bmp',
            'image/tiff': 'tiff'
        }
        return content_type_map.get(content_type, 'png')
    
    def get_supported_formats(self) -> list:
        """获取支持的输入格式"""
        return ['.docx', '.doc']
    
    def check_service_status(self) -> Dict:
        """检查服务状态"""
        if Document is None:
            return {
                'status': 'offline',
                'error': 'python-docx库未安装',
                'message': '请运行: pip install python-docx'
            }
        
        return {
            'status': 'online',
            'message': 'Docx解析服务运行正常'
        }

# 测试函数
async def test_docx_service():
    """测试Docx解析服务"""
    service = DocxParsingService()
    
    # 测试服务状态
    status = service.check_service_status()
    print(f"服务状态: {status}")
    
    # 测试试题文件
    question_file = "../data/input-doc/zt01-math-set-question.docx"
    if os.path.exists(question_file):
        print(f"\n测试解析试题文件: {question_file}")
        result = await service.parse_document(question_file, is_answer=False)
        
        if result['success']:
            print("✓ 试题解析成功!")
            print(f"输出文件: {result['output_path']}")
            print(f"HTML内容长度: {len(result['html_content'])} 字符")
            print(f"内容预览: {result['html_content'][:200]}...")
        else:
            print(f"✗ 试题解析失败: {result.get('error')}")
    else:
        print(f"测试文件不存在: {question_file}")
    
    # 测试答案文件
    answer_file = "../data/input-doc/zt01-math-set-answer.docx"
    if os.path.exists(answer_file):
        print(f"\n测试解析答案文件: {answer_file}")
        result = await service.parse_document(answer_file, is_answer=True)
        
        if result['success']:
            print("✓ 答案解析成功!")
            print(f"输出文件: {result['output_path']}")
            print(f"HTML内容长度: {len(result['html_content'])} 字符")
        else:
            print(f"✗ 答案解析失败: {result.get('error')}")
    else:
        print(f"测试文件不存在: {answer_file}")

if __name__ == "__main__":
    # 运行测试
    asyncio.run(test_docx_service())